"""
Module pour extraire le texte de différents formats de documents.
Supporte : PDF, DOCX, XLSX, TXT
"""

import os
from typing import Optional
import tempfile

# Imports conditionnels pour gérer les dépendances
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None


class DocumentProcessor:
    """Classe pour extraire le texte de différents formats de documents."""

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extrait le texte d'un fichier PDF."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 n'est pas installé. Installez-le avec: pip install PyPDF2")

        text_parts = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du PDF : {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extrait le texte d'un fichier DOCX."""
        if Document is None:
            raise ImportError("python-docx n'est pas installé. Installez-le avec: pip install python-docx")

        try:
            doc = Document(file_path)
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extraction du texte des tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du DOCX : {str(e)}")

    @staticmethod
    def extract_text_from_xlsx(file_path: str) -> str:
        """Extrait le texte d'un fichier XLSX."""
        if load_workbook is None:
            raise ImportError("openpyxl n'est pas installé. Installez-le avec: pip install openpyxl")

        try:
            workbook = load_workbook(file_path, data_only=True)
            text_parts = []

            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_parts.append(f"--- Feuille: {sheet_name} ---")

                for row in sheet.iter_rows(values_only=True):
                    # Filtrer les lignes vides
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    if any(val.strip() for val in row_values):
                        row_text = " | ".join(row_values)
                        text_parts.append(row_text)

            return "\n\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du XLSX : {str(e)}")

    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extrait le texte d'un fichier TXT."""
        try:
            # Tenter différents encodages
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue

            raise Exception("Impossible de décoder le fichier avec les encodages supportés")
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction du TXT : {str(e)}")

    @staticmethod
    def process_uploaded_file(uploaded_file, file_extension: str) -> Optional[str]:
        """
        Traite un fichier uploadé et extrait son texte.

        Args:
            uploaded_file: Fichier uploadé depuis Streamlit
            file_extension: Extension du fichier (avec le point, ex: '.pdf')

        Returns:
            Le texte extrait ou None en cas d'erreur
        """
        # Créer un fichier temporaire
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name

        try:
            if file_extension.lower() == '.pdf':
                text = DocumentProcessor.extract_text_from_pdf(tmp_file_path)
            elif file_extension.lower() == '.docx':
                text = DocumentProcessor.extract_text_from_docx(tmp_file_path)
            elif file_extension.lower() == '.xlsx':
                text = DocumentProcessor.extract_text_from_xlsx(tmp_file_path)
            elif file_extension.lower() == '.txt':
                text = DocumentProcessor.extract_text_from_txt(tmp_file_path)
            else:
                raise ValueError(f"Format de fichier non supporté : {file_extension}")

            return text
        finally:
            # Supprimer le fichier temporaire
            if os.path.exists(tmp_file_path):
                os.remove(tmp_file_path)

    @staticmethod
    def get_supported_extensions() -> list:
        """Retourne la liste des extensions de fichiers supportées."""
        return ['.pdf', '.docx', '.xlsx', '.txt']
